import base64
import json
import mimetypes
import time
from io import BytesIO
from typing import Any

from ai_workflow_mapper.platform.contracts.document_loader import (
    DocumentLoaderConfig,
    DocumentLoaderError,
    DocumentLoaderErrorCode,
    DocumentLoaderOperation,
    DocumentLoaderRequest,
    DocumentLoaderResponse,
    DocumentLoaderStatus,
)

_SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".pdf", ".docx"}


class DocumentLoaderModuleError(Exception):
    def __init__(self, error: DocumentLoaderError) -> None:
        self.error = error
        super().__init__(error.message)


class LocalDocumentLoader:
    """Project-local implementation of the document_loader contract."""

    MODULE_ID = "document_loader"
    DEFAULT_MAX_BYTES = 10 * 1024 * 1024  # 10 MB

    def __init__(self, config: DocumentLoaderConfig) -> None:
        self._config = config
        self._max_bytes: int = config.settings.get("max_size_bytes", self.DEFAULT_MAX_BYTES)
        self._supported: set[str] = set(
            config.settings.get("supported_extensions", list(_SUPPORTED_EXTENSIONS))
        )

    def handle(self, request: DocumentLoaderRequest) -> DocumentLoaderResponse:
        t0 = time.monotonic()
        try:
            result, warnings = self._dispatch(request)
            return DocumentLoaderResponse(
                module_id=self.MODULE_ID,
                operation=request.operation,
                status=DocumentLoaderStatus.succeeded,
                result=result,
                warnings=warnings,
                metadata=self._make_metadata(t0),
                trace_id=request.trace_id,
            )
        except DocumentLoaderModuleError as exc:
            return DocumentLoaderResponse(
                module_id=self.MODULE_ID,
                operation=request.operation,
                status=DocumentLoaderStatus.failed,
                result={"error": exc.error.model_dump(exclude_none=True)},
                warnings=[],
                metadata=self._make_metadata(t0),
                trace_id=request.trace_id,
            )

    def get_config(self) -> DocumentLoaderConfig:
        return self._config

    # ------------------------------------------------------------------
    # Dispatch
    # ------------------------------------------------------------------

    def _dispatch(self, request: DocumentLoaderRequest) -> tuple[dict[str, Any], list[str]]:
        op = request.operation
        if op == DocumentLoaderOperation.detect_file_type:
            return self._detect_file_type(request.input, request.trace_id), []
        if op == DocumentLoaderOperation.load_document:
            return self._load_document(request.input, request.trace_id)
        if op == DocumentLoaderOperation.extract_text:
            return self._extract_text(request.input, request.trace_id)
        if op == DocumentLoaderOperation.extract_metadata:
            return self._extract_metadata(request.input, request.trace_id)
        raise ValueError(f"Unknown operation: {op}")

    # ------------------------------------------------------------------
    # Operations
    # ------------------------------------------------------------------

    def _detect_file_type(self, inp: dict[str, Any], trace_id: str) -> dict[str, Any]:
        filename: str = inp.get("filename", "")
        ext = self._extension(filename)
        mime_type, _ = mimetypes.guess_type(filename)
        return {
            "mime_type": mime_type or "application/octet-stream",
            "extension": ext,
            "supported": ext in self._supported,
        }

    def _load_document(
        self, inp: dict[str, Any], trace_id: str
    ) -> tuple[dict[str, Any], list[str]]:
        filename: str = inp.get("filename", "")
        ext = self._extension(filename)
        if ext not in self._supported:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=DocumentLoaderOperation.load_document,
                    error_code=DocumentLoaderErrorCode.document_unsupported_type,
                    message=f"Unsupported file type: '{ext}'",
                    retryable=False,
                    trace_id=trace_id,
                )
            )

        raw_bytes = self._decode_bytes(inp, DocumentLoaderOperation.load_document, trace_id)
        if len(raw_bytes) > self._max_bytes:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=DocumentLoaderOperation.load_document,
                    error_code=DocumentLoaderErrorCode.document_too_large,
                    message=(
                        f"File size {len(raw_bytes)} bytes exceeds limit of {self._max_bytes} bytes"
                    ),
                    retryable=False,
                    trace_id=trace_id,
                )
            )

        mime_type, _ = mimetypes.guess_type(filename)
        warnings: list[str] = []
        return (
            {
                "raw_bytes_b64": base64.b64encode(raw_bytes).decode(),
                "size_bytes": len(raw_bytes),
                "filename": filename,
                "mime_type": mime_type or "application/octet-stream",
            },
            warnings,
        )

    def _extract_text(
        self, inp: dict[str, Any], trace_id: str
    ) -> tuple[dict[str, Any], list[str]]:
        filename: str = inp.get("filename", "")
        ext = self._extension(filename)
        raw_bytes = self._decode_bytes(inp, DocumentLoaderOperation.extract_text, trace_id)
        warnings: list[str] = []

        if ext in {".txt", ".md"}:
            try:
                text = raw_bytes.decode("utf-8")
            except UnicodeDecodeError as exc:
                raise DocumentLoaderModuleError(
                    DocumentLoaderError(
                        operation=DocumentLoaderOperation.extract_text,
                        error_code=DocumentLoaderErrorCode.document_parse_failed,
                        message=f"UTF-8 decode failed: {exc}",
                        retryable=False,
                        trace_id=trace_id,
                    )
                ) from exc
            return {"text": text, "char_count": len(text), "parser": "plaintext"}, warnings

        if ext == ".json":
            try:
                obj = json.loads(raw_bytes.decode("utf-8"))
                text = json.dumps(obj, indent=2, ensure_ascii=False)
            except (json.JSONDecodeError, UnicodeDecodeError) as exc:
                raise DocumentLoaderModuleError(
                    DocumentLoaderError(
                        operation=DocumentLoaderOperation.extract_text,
                        error_code=DocumentLoaderErrorCode.document_parse_failed,
                        message=f"JSON parse failed: {exc}",
                        retryable=False,
                        trace_id=trace_id,
                    )
                ) from exc
            return {"text": text, "char_count": len(text), "parser": "json"}, warnings

        if ext == ".pdf":
            return self._extract_text_pdf(raw_bytes, trace_id)

        if ext == ".docx":
            return self._extract_text_docx(raw_bytes, trace_id)

        raise DocumentLoaderModuleError(
            DocumentLoaderError(
                operation=DocumentLoaderOperation.extract_text,
                error_code=DocumentLoaderErrorCode.document_unsupported_type,
                message=f"Unsupported file type for text extraction: '{ext}'",
                retryable=False,
                trace_id=trace_id,
            )
        )

    def _extract_text_pdf(
        self, raw_bytes: bytes, trace_id: str
    ) -> tuple[dict[str, Any], list[str]]:
        try:
            import pypdf  # noqa: PLC0415
        except ImportError as exc:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=DocumentLoaderOperation.extract_text,
                    error_code=DocumentLoaderErrorCode.document_parse_failed,
                    message="pypdf is not installed",
                    retryable=False,
                    trace_id=trace_id,
                )
            ) from exc

        try:
            reader = pypdf.PdfReader(BytesIO(raw_bytes))
        except Exception as exc:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=DocumentLoaderOperation.extract_text,
                    error_code=DocumentLoaderErrorCode.document_parse_failed,
                    message=f"PDF parse failed: {exc}",
                    retryable=False,
                    trace_id=trace_id,
                )
            ) from exc

        if reader.is_encrypted:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=DocumentLoaderOperation.extract_text,
                    error_code=DocumentLoaderErrorCode.document_password_required,
                    message="PDF is encrypted and requires a password",
                    retryable=False,
                    trace_id=trace_id,
                )
            )

        warnings: list[str] = []
        pages: list[str] = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if not page_text.strip():
                warnings.append(f"Page {i + 1} yielded no text (may be image-based)")
            pages.append(page_text)

        text = "\f".join(pages)
        return {
            "text": text,
            "char_count": len(text),
            "page_count": len(reader.pages),
            "parser": "pypdf",
        }, warnings

    def _extract_text_docx(
        self, raw_bytes: bytes, trace_id: str
    ) -> tuple[dict[str, Any], list[str]]:
        try:
            import docx  # noqa: PLC0415
        except ImportError as exc:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=DocumentLoaderOperation.extract_text,
                    error_code=DocumentLoaderErrorCode.document_parse_failed,
                    message="python-docx is not installed",
                    retryable=False,
                    trace_id=trace_id,
                )
            ) from exc

        try:
            doc = docx.Document(BytesIO(raw_bytes))
        except Exception as exc:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=DocumentLoaderOperation.extract_text,
                    error_code=DocumentLoaderErrorCode.document_parse_failed,
                    message=f"DOCX parse failed: {exc}",
                    retryable=False,
                    trace_id=trace_id,
                )
            ) from exc

        _W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        def _cell_text(cell_el: Any) -> str:
            return "".join(n.text or "" for n in cell_el.iter(f"{{{_W}}}t")).strip()

        # Headers (prepended; skip sections linked to previous to avoid duplicates)
        header_lines: list[str] = []
        for section in doc.sections:
            if not section.header.is_linked_to_previous:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_lines.append(para.text)

        # Body in document order — paragraphs and tables interleaved
        body_blocks: list[str] = []
        for element in doc.element.body:
            local = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            if local == "p":
                text = "".join(n.text or "" for n in element.iter(f"{{{_W}}}t"))
                if text.strip():
                    body_blocks.append(text)
            elif local == "tbl":
                for row in element.findall(f".//{{{_W}}}tr"):
                    cells = [_cell_text(c) for c in row.findall(f"{{{_W}}}tc")]
                    if any(cells):
                        body_blocks.append(" | ".join(cells))

        # Footers (appended; same linked-section logic)
        footer_lines: list[str] = []
        for section in doc.sections:
            if not section.footer.is_linked_to_previous:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_lines.append(para.text)

        parts: list[str] = []
        if header_lines:
            parts.append("\n".join(header_lines))
        parts.extend(body_blocks)
        if footer_lines:
            parts.append("\n".join(footer_lines))
        text = "\n".join(parts)

        return {"text": text, "char_count": len(text), "parser": "python-docx"}, []

    def _extract_metadata(
        self, inp: dict[str, Any], trace_id: str
    ) -> tuple[dict[str, Any], list[str]]:
        filename: str = inp.get("filename", "")
        ext = self._extension(filename)
        raw_bytes = self._decode_bytes(inp, DocumentLoaderOperation.extract_metadata, trace_id)
        mime_type, _ = mimetypes.guess_type(filename)

        meta: dict[str, Any] = {
            "filename": filename,
            "extension": ext,
            "size_bytes": len(raw_bytes),
            "mime_type": mime_type or "application/octet-stream",
        }

        if ext == ".json":
            try:
                obj = json.loads(raw_bytes.decode("utf-8"))
                meta["json_keys_top_level"] = list(obj.keys()) if isinstance(obj, dict) else []
            except Exception:
                pass

        elif ext == ".pdf":
            try:
                import pypdf  # noqa: PLC0415

                reader = pypdf.PdfReader(BytesIO(raw_bytes))
                pdf_meta = reader.metadata or {}
                meta["page_count"] = len(reader.pages)
                meta["is_encrypted"] = reader.is_encrypted
                meta["title"] = pdf_meta.get("/Title", "")
                meta["author"] = pdf_meta.get("/Author", "")
            except Exception:
                pass

        elif ext == ".docx":
            try:
                import docx  # noqa: PLC0415

                doc = docx.Document(BytesIO(raw_bytes))
                props = doc.core_properties
                meta["paragraph_count"] = len(doc.paragraphs)
                meta["table_count"] = len(doc.tables)
                meta["core_properties"] = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "created": props.created.isoformat() if props.created else None,
                }
            except Exception:
                pass

        return meta, []

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _extension(self, filename: str) -> str:
        import os

        _, ext = os.path.splitext(filename.lower())
        return ext

    def _decode_bytes(
        self,
        inp: dict[str, Any],
        operation: DocumentLoaderOperation,
        trace_id: str,
    ) -> bytes:
        b64 = inp.get("content_bytes_b64", "")
        if not b64:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=operation,
                    error_code=DocumentLoaderErrorCode.document_parse_failed,
                    message="Missing 'content_bytes_b64' in input",
                    retryable=False,
                    trace_id=trace_id,
                )
            )
        try:
            return base64.b64decode(b64)
        except Exception as exc:
            raise DocumentLoaderModuleError(
                DocumentLoaderError(
                    operation=operation,
                    error_code=DocumentLoaderErrorCode.document_parse_failed,
                    message=f"Invalid base64 content: {exc}",
                    retryable=False,
                    trace_id=trace_id,
                )
            ) from exc

    def _make_metadata(self, t0: float) -> dict[str, Any]:
        return {
            "implementation": "local",
            "contract_version": "0.1.0",
            "latency_ms": round((time.monotonic() - t0) * 1000, 2),
        }
