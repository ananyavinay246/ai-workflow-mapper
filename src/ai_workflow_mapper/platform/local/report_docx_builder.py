"""Build DOCX analysis reports from validated structured data."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt


def build_report_docx(data: dict[str, Any], metadata: dict[str, Any] | None = None) -> bytes:
    """Return DOCX bytes mirroring the seven report.md.j2 sections."""
    meta = metadata or {}
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    doc.add_heading("Workflow Analysis Report", level=0)
    doc.add_paragraph(f"Generated: {meta.get('generated_at', 'N/A')}")
    doc.add_paragraph(f"Job ID: {meta.get('job_id', 'N/A')}")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(data.get("executive_summary") or "No summary provided.")

    doc.add_heading("Process Inventory", level=1)
    processes = data.get("processes") or []
    if processes:
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        for i, label in enumerate(["#", "Process Name", "Owner", "Duration", "Frequency"]):
            hdr[i].text = label
        for idx, proc in enumerate(processes, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = str(proc.get("name", ""))
            row[2].text = str(proc.get("owner") or "—")
            row[3].text = str(proc.get("duration") or "—")
            row[4].text = str(proc.get("frequency") or "—")
    else:
        doc.add_paragraph("No processes identified.")

    doc.add_heading("Bottleneck Analysis", level=1)
    bottlenecks = data.get("bottlenecks") or []
    if bottlenecks:
        for bn in bottlenecks:
            doc.add_heading(
                f"{bn.get('name', 'Bottleneck')} — Severity: {bn.get('severity', 'Unknown')}",
                level=2,
            )
            doc.add_paragraph(bn.get("description") or "")
            if bn.get("impact"):
                doc.add_paragraph(f"Impact: {bn['impact']}")
            if bn.get("root_cause_hypothesis"):
                doc.add_paragraph(f"Root cause hypothesis: {bn['root_cause_hypothesis']}")
    else:
        doc.add_paragraph("No bottlenecks detected.")

    doc.add_heading("Redundancy Analysis", level=1)
    redundancies = data.get("redundancies") or []
    if redundancies:
        for rd in redundancies:
            doc.add_heading(str(rd.get("name", "Redundancy")), level=2)
            doc.add_paragraph(rd.get("description") or "")
            if rd.get("waste_estimate"):
                doc.add_paragraph(f"Waste estimate: {rd['waste_estimate']}")
            affected = rd.get("affected_steps") or []
            if affected:
                doc.add_paragraph(f"Affected steps: {', '.join(affected)}")
    else:
        doc.add_paragraph("No redundancies detected.")

    doc.add_heading("Automation Opportunity Matrix", level=1)
    opportunities = data.get("automation_opportunities") or []
    if opportunities:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        for i, label in enumerate(["Opportunity", "Effort", "ROI", "Priority"]):
            hdr[i].text = label
        for ao in opportunities:
            row = table.add_row().cells
            row[0].text = str(ao.get("name", ""))
            row[1].text = str(ao.get("effort") or "—")
            row[2].text = str(ao.get("roi") or "—")
            row[3].text = str(ao.get("priority") or "—")
            if ao.get("suggested_approach"):
                doc.add_paragraph(f"Suggested approach ({ao.get('name')}): {ao['suggested_approach']}")
            if ao.get("time_savings_per_week"):
                doc.add_paragraph(f"Time savings ({ao.get('name')}): {ao['time_savings_per_week']}")
    else:
        doc.add_paragraph("No automation opportunities identified.")

    doc.add_heading("Recommended Next Steps", level=1)
    next_steps = data.get("next_steps") or []
    if next_steps:
        for i, step in enumerate(next_steps, start=1):
            doc.add_paragraph(f"{i}. {step}")
    else:
        doc.add_paragraph("No recommendations available.")

    skipped = meta.get("skipped_documents")
    if skipped is not None:
        doc.add_paragraph(f"Documents skipped during normalization: {skipped}")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
