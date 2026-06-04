"""Contract tests for LocalDocumentLoader (document_loader module)."""

import base64
import json
from io import BytesIO
from pathlib import Path

import jsonschema
import pytest

from ai_workflow_mapper.platform.contracts.document_loader import (
    DocumentLoaderConfig,
    DocumentLoaderContext,
    DocumentLoaderOperation,
    DocumentLoaderRequest,
)
from ai_workflow_mapper.platform.local.document_loader import LocalDocumentLoader

SCHEMAS_DIR = (
    Path(__file__).parents[2] / "shared_modules" / "document_loader" / "schemas"
)


def _load_schema(name: str) -> dict:
    return json.loads((SCHEMAS_DIR / name).read_text())


def _make_loader(extra_settings: dict | None = None) -> LocalDocumentLoader:
    settings: dict = {"max_size_bytes": 1024 * 1024}
    if extra_settings:
        settings.update(extra_settings)
    config = DocumentLoaderConfig(
        environment="local",
        implementation="local",
        settings=settings,
        security={},
    )
    return LocalDocumentLoader(config)


def _make_request(operation: DocumentLoaderOperation, inp: dict) -> DocumentLoaderRequest:
    return DocumentLoaderRequest(
        request_id="test-001",
        operation=operation,
        input=inp,
        context=DocumentLoaderContext(actor_id="test", tenant_id="test", environment="local"),
        trace_id="trace-test-001",
    )


def _b64(content: str) -> str:
    return base64.b64encode(content.encode()).decode()


def _b64_bytes(content: bytes) -> str:
    return base64.b64encode(content).decode()


# ---------------------------------------------------------------------------
# detect_file_type
# ---------------------------------------------------------------------------


def test_detect_file_type_txt():
    loader = _make_loader()
    resp = loader.handle(_make_request(DocumentLoaderOperation.detect_file_type, {"filename": "process.txt"}))
    assert resp.status.value == "succeeded"
    assert resp.result["extension"] == ".txt"
    assert resp.result["supported"] is True


def test_detect_file_type_md():
    loader = _make_loader()
    resp = loader.handle(_make_request(DocumentLoaderOperation.detect_file_type, {"filename": "notes.md"}))
    assert resp.status.value == "succeeded"
    assert resp.result["supported"] is True


def test_detect_file_type_pdf():
    loader = _make_loader()
    resp = loader.handle(_make_request(DocumentLoaderOperation.detect_file_type, {"filename": "report.pdf"}))
    assert resp.status.value == "succeeded"
    assert resp.result["supported"] is True
    assert "pdf" in resp.result["mime_type"]


def test_detect_file_type_docx():
    loader = _make_loader()
    resp = loader.handle(_make_request(DocumentLoaderOperation.detect_file_type, {"filename": "doc.docx"}))
    assert resp.status.value == "succeeded"
    assert resp.result["supported"] is True


def test_detect_file_type_unsupported():
    loader = _make_loader()
    resp = loader.handle(_make_request(DocumentLoaderOperation.detect_file_type, {"filename": "virus.exe"}))
    assert resp.status.value == "succeeded"
    assert resp.result["supported"] is False


# ---------------------------------------------------------------------------
# load_document
# ---------------------------------------------------------------------------


def test_load_document_txt_happy():
    loader = _make_loader()
    content = "Hello, workflow!"
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.load_document,
            {"filename": "process.txt", "content_bytes_b64": _b64(content)},
        )
    )
    assert resp.status.value == "succeeded"
    assert resp.result["size_bytes"] == len(content.encode())
    assert resp.result["filename"] == "process.txt"


def test_load_document_too_large():
    loader = _make_loader({"max_size_bytes": 10})
    content = "A" * 20
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.load_document,
            {"filename": "big.txt", "content_bytes_b64": _b64(content)},
        )
    )
    assert resp.status.value == "failed"
    assert resp.result["error"]["error_code"] == "document_too_large"
    assert resp.result["error"]["retryable"] is False


def test_load_document_unsupported_type():
    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.load_document,
            {"filename": "file.exe", "content_bytes_b64": _b64("data")},
        )
    )
    assert resp.status.value == "failed"
    assert resp.result["error"]["error_code"] == "document_unsupported_type"


def test_load_document_missing_content():
    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.load_document,
            {"filename": "process.txt"},
        )
    )
    assert resp.status.value == "failed"
    assert resp.result["error"]["error_code"] == "document_parse_failed"


# ---------------------------------------------------------------------------
# extract_text
# ---------------------------------------------------------------------------


def test_extract_text_txt():
    loader = _make_loader()
    text = "Step 1: Submit request\nStep 2: Approval"
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "process.txt", "content_bytes_b64": _b64(text)},
        )
    )
    assert resp.status.value == "succeeded"
    assert resp.result["text"] == text
    assert resp.result["parser"] == "plaintext"
    assert resp.result["char_count"] == len(text)


def test_extract_text_md():
    loader = _make_loader()
    text = "# Workflow\n\n- Step 1\n- Step 2"
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "workflow.md", "content_bytes_b64": _b64(text)},
        )
    )
    assert resp.status.value == "succeeded"
    assert resp.result["parser"] == "plaintext"


def test_extract_text_json():
    loader = _make_loader()
    obj = {"steps": ["approval", "review"], "owner": "alice"}
    raw = json.dumps(obj)
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "data.json", "content_bytes_b64": _b64(raw)},
        )
    )
    assert resp.status.value == "succeeded"
    assert resp.result["parser"] == "json"
    # Verify round-trip: the text is valid JSON
    parsed = json.loads(resp.result["text"])
    assert parsed["steps"] == ["approval", "review"]


def test_extract_text_json_malformed():
    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "bad.json", "content_bytes_b64": _b64("{not json}")},
        )
    )
    assert resp.status.value == "failed"
    assert resp.result["error"]["error_code"] == "document_parse_failed"


def test_extract_text_pdf_encrypted(tmp_path):
    """Encrypted PDFs return document_password_required."""
    try:
        import pypdf
    except ImportError:
        pytest.skip("pypdf not installed")

    # Build a minimal encrypted PDF using pypdf's writer
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.encrypt("secret")
    buf = BytesIO()
    writer.write(buf)
    raw = buf.getvalue()

    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "locked.pdf", "content_bytes_b64": _b64_bytes(raw)},
        )
    )
    assert resp.status.value == "failed"
    assert resp.result["error"]["error_code"] == "document_password_required"


def test_extract_text_pdf_happy(tmp_path):
    """A simple single-page PDF returns text and page_count."""
    try:
        import pypdf
    except ImportError:
        pytest.skip("pypdf not installed")

    from io import BytesIO

    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = BytesIO()
    writer.write(buf)
    raw = buf.getvalue()

    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "blank.pdf", "content_bytes_b64": _b64_bytes(raw)},
        )
    )
    assert resp.status.value == "succeeded"
    assert resp.result["parser"] == "pypdf"
    assert resp.result["page_count"] == 1


def test_extract_text_docx_happy():
    """A minimal DOCX document returns text and parser=python-docx."""
    try:
        import docx
        from io import BytesIO
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = docx.Document()
    doc.add_paragraph("Approve request")
    doc.add_paragraph("Send to manager")
    buf = BytesIO()
    doc.save(buf)
    raw = buf.getvalue()

    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "workflow.docx", "content_bytes_b64": _b64_bytes(raw)},
        )
    )
    assert resp.status.value == "succeeded"
    assert "Approve request" in resp.result["text"]
    assert resp.result["parser"] == "python-docx"


def test_extract_text_docx_table():
    """Table cell text is extracted and rows are pipe-separated."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "step"
    table.cell(0, 1).text = "owner"
    table.cell(1, 0).text = "Submit"
    table.cell(1, 1).text = "HR"
    buf = BytesIO()
    doc.save(buf)

    resp = _make_loader().handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "sop.docx", "content_bytes_b64": _b64_bytes(buf.getvalue())},
        )
    )
    assert resp.status.value == "succeeded"
    assert "Submit" in resp.result["text"]
    assert "HR" in resp.result["text"]
    assert "Submit | HR" in resp.result["text"]


def test_extract_text_docx_header_footer():
    """Header and footer paragraphs appear in extracted text."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = docx.Document()
    doc.add_paragraph("Body paragraph")
    section = doc.sections[0]
    section.header.paragraphs[0].text = "CONFIDENTIAL"
    section.footer.paragraphs[0].text = "Page 1"
    buf = BytesIO()
    doc.save(buf)

    resp = _make_loader().handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "sop.docx", "content_bytes_b64": _b64_bytes(buf.getvalue())},
        )
    )
    assert resp.status.value == "succeeded"
    text = resp.result["text"]
    assert "CONFIDENTIAL" in text
    assert "Page 1" in text
    assert "Body paragraph" in text


def test_extract_text_docx_body_order():
    """Paragraph before a table appears before the table row in extracted text."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = docx.Document()
    doc.add_paragraph("Intro paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Step"
    table.cell(0, 1).text = "Owner"
    doc.add_paragraph("Closing paragraph")
    buf = BytesIO()
    doc.save(buf)

    resp = _make_loader().handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "sop.docx", "content_bytes_b64": _b64_bytes(buf.getvalue())},
        )
    )
    assert resp.status.value == "succeeded"
    text = resp.result["text"]
    assert text.index("Intro paragraph") < text.index("Step")
    assert text.index("Step") < text.index("Closing paragraph")


def test_extract_text_docx_empty_cells_skipped():
    """Rows with all-empty cells are omitted; rows with some empty cells still appear."""
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Submit"
    table.cell(0, 1).text = ""   # empty second cell
    # row 1: both cells empty — should be omitted
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = ""
    buf = BytesIO()
    doc.save(buf)

    resp = _make_loader().handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "sop.docx", "content_bytes_b64": _b64_bytes(buf.getvalue())},
        )
    )
    assert resp.status.value == "succeeded"
    text = resp.result["text"]
    assert "Submit" in text
    # only one row line should appear (the all-empty row is dropped)
    pipe_lines = [ln for ln in text.splitlines() if "|" in ln]
    assert len(pipe_lines) == 1


# ---------------------------------------------------------------------------
# extract_metadata
# ---------------------------------------------------------------------------


def test_extract_metadata_txt():
    loader = _make_loader()
    content = "Some process text"
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_metadata,
            {"filename": "process.txt", "content_bytes_b64": _b64(content)},
        )
    )
    assert resp.status.value == "succeeded"
    assert resp.result["extension"] == ".txt"
    assert resp.result["size_bytes"] == len(content.encode())
    assert resp.result["filename"] == "process.txt"


def test_extract_metadata_json():
    loader = _make_loader()
    obj = {"key1": 1, "key2": 2}
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_metadata,
            {"filename": "data.json", "content_bytes_b64": _b64(json.dumps(obj))},
        )
    )
    assert resp.status.value == "succeeded"
    assert set(resp.result["json_keys_top_level"]) == {"key1", "key2"}


def test_extract_metadata_pdf():
    try:
        import pypdf
        from io import BytesIO
    except ImportError:
        pytest.skip("pypdf not installed")

    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = BytesIO()
    writer.write(buf)
    raw = buf.getvalue()

    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_metadata,
            {"filename": "report.pdf", "content_bytes_b64": _b64_bytes(raw)},
        )
    )
    assert resp.status.value == "succeeded"
    assert "page_count" in resp.result


# ---------------------------------------------------------------------------
# Schema conformance
# ---------------------------------------------------------------------------


def test_input_validates_against_schema():
    schema = _load_schema("input.schema.json")
    req = _make_request(DocumentLoaderOperation.detect_file_type, {"filename": "test.txt"})
    jsonschema.validate(req.model_dump(), schema)


def test_output_validates_against_schema():
    schema = _load_schema("output.schema.json")
    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "workflow.txt", "content_bytes_b64": _b64("hello")},
        )
    )
    jsonschema.validate(resp.model_dump(), schema)


def test_error_validates_against_schema():
    schema = _load_schema("error.schema.json")
    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.load_document,
            {"filename": "bad.exe", "content_bytes_b64": _b64("data")},
        )
    )
    assert resp.status.value == "failed"
    jsonschema.validate(resp.result["error"], schema)


def test_config_validates_against_schema():
    schema = _load_schema("config.schema.json")
    loader = _make_loader()
    jsonschema.validate(loader.get_config().model_dump(exclude_none=True), schema)


# ---------------------------------------------------------------------------
# Security: no sensitive values in metadata
# ---------------------------------------------------------------------------


def test_no_sensitive_values_in_metadata():
    loader = _make_loader()
    content = "SECRET_KEY=abc123"
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.extract_text,
            {"filename": "config.txt", "content_bytes_b64": _b64(content)},
        )
    )
    # metadata must not contain raw file content
    meta_str = json.dumps(resp.metadata)
    assert "SECRET_KEY" not in meta_str
    assert "abc123" not in meta_str


# ---------------------------------------------------------------------------
# Backward compatibility: required fields present
# ---------------------------------------------------------------------------


def test_backward_compatible_fields_stable():
    loader = _make_loader()
    resp = loader.handle(
        _make_request(
            DocumentLoaderOperation.detect_file_type,
            {"filename": "test.txt"},
        )
    )
    for field in ("module_id", "operation", "status", "result", "warnings", "metadata", "trace_id"):
        assert hasattr(resp, field), f"Missing field: {field}"
    assert resp.module_id == "document_loader"
    assert resp.trace_id == "trace-test-001"
