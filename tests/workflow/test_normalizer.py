"""Tests for InputNormalizer — all input format tiers."""

import base64
import json
import zipfile
from io import BytesIO

import pytest

from ai_workflow_mapper.platform.contracts.document_loader import DocumentLoaderConfig
from ai_workflow_mapper.platform.local.document_loader import LocalDocumentLoader
from ai_workflow_mapper.workflow.domain import InputDocument, WorkflowInput
from ai_workflow_mapper.workflow.normalizer import InputNormalizer


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _b64(text: str) -> str:
    return base64.b64encode(text.encode()).decode()


def _b64_bytes(raw: bytes) -> str:
    return base64.b64encode(raw).decode()


def _make_loader() -> LocalDocumentLoader:
    return LocalDocumentLoader(
        DocumentLoaderConfig(
            environment="local",
            implementation="local",
            settings={},
            security={},
        )
    )


def _make_normalizer() -> InputNormalizer:
    return InputNormalizer(_make_loader())


def _make_doc(filename: str, content: str | bytes, source_type: str = "document") -> InputDocument:
    if isinstance(content, str):
        b64 = _b64(content)
    else:
        b64 = _b64_bytes(content)
    return InputDocument(filename=filename, content_b64=b64, source_type=source_type)


def _make_input(*docs: InputDocument, description: str | None = None) -> WorkflowInput:
    return WorkflowInput(documents=list(docs), description=description)


def _norm(normalizer: InputNormalizer, *docs: InputDocument, description: str | None = None):
    return normalizer.normalize(_make_input(*docs, description=description), trace_id="test-trace")


def _minimal_vsdx(page_texts: list[str]) -> bytes:
    """Build a minimal in-memory VSDX (ZIP) containing page XML files."""
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        # [Content_Types].xml — minimal required entry
        zf.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="utf-8"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>'
        )
        for i, text in enumerate(page_texts, start=1):
            xml = (
                '<?xml version="1.0" encoding="utf-8"?>'
                '<VisioDocument xmlns="http://schemas.microsoft.com/office/visio/2012/main">'
                f"<Page><Shapes><Shape><Text>{text}</Text></Shape></Shapes></Page>"
                "</VisioDocument>"
            )
            zf.writestr(f"visio/pages/page{i}.xml", xml)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Tier 1 — document_loader formats
# ---------------------------------------------------------------------------


def test_normalize_txt():
    n = _make_normalizer()
    result = _norm(n, _make_doc("process.txt", "Step 1: Submit\nStep 2: Approve"))
    assert len(result.documents) == 1
    doc = result.documents[0]
    assert doc.parser == "plaintext"
    assert "Step 1" in doc.text
    assert doc.char_count > 0
    assert doc.filename == "process.txt"


def test_normalize_md():
    n = _make_normalizer()
    result = _norm(n, _make_doc("runbook.md", "# Workflow\n- Step A\n- Step B"))
    assert len(result.documents) == 1
    assert result.documents[0].parser == "plaintext"
    assert "Step A" in result.documents[0].text


def test_normalize_json():
    obj = {"steps": ["Approve", "Review"], "owner": "ops"}
    result = _norm(_make_normalizer(), _make_doc("export.json", json.dumps(obj)))
    assert len(result.documents) == 1
    doc = result.documents[0]
    assert doc.parser == "json"
    parsed = json.loads(doc.text)
    assert parsed["steps"] == ["Approve", "Review"]


def test_normalize_pdf():
    try:
        import pypdf
    except ImportError:
        pytest.skip("pypdf not installed")

    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = BytesIO()
    writer.write(buf)
    raw = buf.getvalue()

    result = _norm(_make_normalizer(), _make_doc("report.pdf", raw))
    assert len(result.documents) == 1
    doc = result.documents[0]
    assert doc.parser == "pypdf"
    assert doc.metadata.get("page_count") == 1


def test_normalize_docx():
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx not installed")

    d = docx.Document()
    d.add_paragraph("Approve the request")
    d.add_paragraph("Send to manager")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "step"
    table.cell(0, 1).text = "owner"
    table.cell(1, 0).text = "Submit"
    table.cell(1, 1).text = "HR"
    buf = BytesIO()
    d.save(buf)
    raw = buf.getvalue()

    result = _norm(_make_normalizer(), _make_doc("sop.docx", raw))
    assert len(result.documents) == 1
    doc = result.documents[0]
    assert doc.parser == "python-docx"
    assert "Approve" in doc.text
    assert "Submit" in doc.text
    assert "HR" in doc.text
    assert "| --- | --- |" in doc.text
    assert doc.metadata.get("cleaning", {}).get("tables_converted", 0) >= 1


def test_normalize_txt_cleaning_footer_and_table():
    raw = (
        "CONFIDENTIAL\n"
        "Process steps\n"
        "step | owner\n"
        "Submit | HR\n"
        "Page 1 of 1\n"
    )
    result = _norm(_make_normalizer(), _make_doc("process.txt", raw))
    assert len(result.documents) == 1
    doc = result.documents[0]
    assert "Page 1 of 1" not in doc.text
    assert "| step | owner |" in doc.text
    assert "| Submit | HR |" in doc.text
    assert doc.char_count == len(doc.text)
    assert "cleaning" in doc.metadata


def test_normalize_json_preserves_structure():
    obj = {"steps": ["Approve", "Review"], "owner": "ops"}
    result = _norm(_make_normalizer(), _make_doc("export.json", json.dumps(obj)))
    doc = result.documents[0]
    parsed = json.loads(doc.text)
    assert parsed["steps"] == ["Approve", "Review"]
    assert "| --- |" not in doc.text


# ---------------------------------------------------------------------------
# Tier 3 — deferred / unsupported → skip
# ---------------------------------------------------------------------------


def test_normalize_url_skipped():
    url_doc = InputDocument(
        filename="https://wiki.internal/sop",
        content_b64=_b64("unused"),
    )
    result = _norm(_make_normalizer(), url_doc)
    assert len(result.documents) == 0
    assert len(result.skipped) == 1
    assert "URL" in result.skipped[0]["reason"]
    assert len(result.warnings) == 1


def test_normalize_unsupported_skipped():
    result = _norm(_make_normalizer(), _make_doc("virus.exe", "binary"))
    assert len(result.documents) == 0
    assert len(result.skipped) == 1
    assert ".exe" in result.skipped[0]["reason"]


# ---------------------------------------------------------------------------
# Failure isolation
# ---------------------------------------------------------------------------


def test_normalize_bad_doc_continues():
    """A corrupt PDF is skipped; a good TXT after it is still normalized."""
    corrupt = _make_doc("broken.pdf", "not a pdf at all")
    good = _make_doc("ok.txt", "Step 1: Review")
    result = _norm(_make_normalizer(), corrupt, good)
    assert len(result.documents) == 1
    assert result.documents[0].filename == "ok.txt"
    assert len(result.skipped) == 1
    assert result.skipped[0]["filename"] == "broken.pdf"


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


def test_normalize_empty_documents():
    result = _norm(_make_normalizer())
    assert result.documents == []
    assert result.skipped == []
    assert result.warnings == []


def test_normalize_multiple_docs():
    docs = [
        _make_doc("a.txt", "alpha"),
        _make_doc("b.txt", "beta"),
        _make_doc("c.txt", "gamma"),
    ]
    result = _norm(_make_normalizer(), *docs)
    assert len(result.documents) == 3
    filenames = [d.filename for d in result.documents]
    assert filenames == ["a.txt", "b.txt", "c.txt"]


def test_normalized_text_not_tagged():
    """Normalizer output must not contain trust wrapper — that is the Extractor's job."""
    result = _norm(_make_normalizer(), _make_doc("proc.txt", "Approve the invoice"))
    assert len(result.documents) == 1
    assert "<process_content" not in result.documents[0].text


def test_normalize_description_only():
    wi = WorkflowInput(documents=[], description="Our onboarding is slow")
    result = _make_normalizer().normalize(wi, trace_id="t")
    assert result.documents == []
    assert result.skipped == []
